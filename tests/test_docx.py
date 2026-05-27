from __future__ import annotations

import json
import base64
import sys
import tempfile
import unittest
import subprocess
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCX_SCRIPTS = ROOT / "skills" / "docx-translator" / "scripts"


def _has_python_docx() -> bool:
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        return False


@unittest.skipUnless(_has_python_docx(), "python-docx 未安装")
class ExportParagraphsTests(unittest.TestCase):
    """export_paragraphs.py 单元测试"""

    def test_creates_output_file(self) -> None:
        """验证导出文件生成和内容格式"""
        with tempfile.TemporaryDirectory() as tmp:
            # 创建最小 .docx
            from docx import Document
            doc = Document()
            doc.add_paragraph("Hello world")
            doc.add_paragraph("Test paragraph 2")
            input_path = Path(tmp) / "test.docx"
            doc.save(str(input_path))

            output_path = Path(tmp) / "paragraphs.txt"
            result = subprocess.run(
                [sys.executable, str(DOCX_SCRIPTS / "export_paragraphs.py"),
                 str(input_path), str(output_path)],
                capture_output=True, text=True, encoding="utf-8",
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertTrue(output_path.exists())
            content = output_path.read_text(encoding="utf-8")
            self.assertIn("Total paragraphs: 2", content)
            self.assertIn("[0]", content)
            self.assertIn("[1]", content)
            self.assertIn("Hello world", content)

    def test_reports_missing_file(self) -> None:
        """不存在的文件应报错退出"""
        result = subprocess.run(
            [sys.executable, str(DOCX_SCRIPTS / "export_paragraphs.py"),
             "nonexistent.docx"],
            capture_output=True, text=True, encoding="utf-8",
        )
        self.assertEqual(result.returncode, 1)


@unittest.skipUnless(_has_python_docx(), "python-docx 未安装")
class ApplyTranslationsTests(unittest.TestCase):
    """apply_translations.py 单元测试"""

    def test_applies_translations(self) -> None:
        """验证翻译文本正确应用到段落"""
        with tempfile.TemporaryDirectory() as tmp:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Original text")
            doc.add_paragraph("Keep this")
            input_path = Path(tmp) / "input.docx"
            doc.save(str(input_path))

            # 创建 base64 编码翻译文件
            translations = {
                "0": base64.b64encode("翻译后的中文文本".encode("utf-8")).decode("ascii"),
            }
            b64_path = Path(tmp) / "translations.json"
            b64_path.write_text(json.dumps(translations, ensure_ascii=False), encoding="utf-8")

            output_path = Path(tmp) / "output.docx"
            result = subprocess.run(
                [sys.executable, str(DOCX_SCRIPTS / "apply_translations.py"),
                 str(input_path), str(b64_path), str(output_path)],
                capture_output=True, text=True, encoding="utf-8",
            )
            self.assertEqual(result.returncode, 0, result.stderr)

            # 验证输出
            out_doc = Document(str(output_path))
            self.assertIn("翻译后的中文文本", out_doc.paragraphs[0].text)
            self.assertIn("Keep this", out_doc.paragraphs[1].text)

    def test_missing_files_report_error(self) -> None:
        """缺失文件应报错"""
        result = subprocess.run(
            [sys.executable, str(DOCX_SCRIPTS / "apply_translations.py"),
             "nonexistent.docx", "nope.json", "out.docx"],
            capture_output=True, text=True, encoding="utf-8",
        )
        self.assertEqual(result.returncode, 1)


if __name__ == "__main__":
    unittest.main()
