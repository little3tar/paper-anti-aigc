#!/usr/bin/env python3
"""Export paragraph list from a .docx file for translation mapping.

Usage:
    uv run --with python-docx python export_paragraphs.py <input.docx> [output.txt]

Output format per line: [index] [style_name] text_content_first_200_chars
"""

import sys, os
from pathlib import Path

SKILL_DIR = Path(__file__).resolve().parent.parent
from docx import Document

def main():
    if len(sys.argv) < 2:
        print("Usage: export_paragraphs.py <input.docx> [output.txt]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'paragraphs.txt'

    if not os.path.exists(input_path):
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    doc = Document(input_path)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"# Total paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"# Total tables: {len(doc.tables)}\n\n")

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                style = para.style.name if para.style else 'None'
                # Truncate long text for readability
                text = para.text[:200]
                f.write(f"[{i}] [{style}] {text}\n")

        # Also list table structure for reference
        if doc.tables:
            f.write("\n# === Table Structures ===\n")
            for ti, table in enumerate(doc.tables):
                f.write(f"# Table {ti}: {len(table.rows)} rows x {len(table.rows[0].cells) if table.rows else 0} cols\n")
                for ri, row in enumerate(table.rows):
                    if ri < 3:  # First 3 rows only
                        cells = [cell.text[:40].replace('\n', '\\n') for cell in row.cells]
                        f.write(f"#   Row {ri}: {' | '.join(cells)}\n")

    print(f"Exported {len(doc.paragraphs)} paragraphs to: {output_path}")
    print(f"Non-empty paragraphs: {sum(1 for p in doc.paragraphs if p.text.strip())}")

if __name__ == '__main__':
    main()
