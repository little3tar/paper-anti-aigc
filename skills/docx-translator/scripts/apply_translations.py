#!/usr/bin/env python3
"""Apply base64-encoded translations to a .docx file.

Reads translations from a JSON file where values are base64-encoded UTF-8 strings,
applies them to matching paragraph indices, and saves a new .docx.

Usage:
    uv run --with python-docx python apply_translations.py <input.docx> <translations_b64.json> <output.docx>

JSON format:
    {"0": "base64encoded...", "5": "base64encoded...", ...}
    Keys are paragraph indices (as strings), values are base64-encoded UTF-8 text.
"""

import json, os, sys, base64
from docx import Document


def main():
    if len(sys.argv) < 4:
        print("Usage: apply_translations.py <input.docx> <translations_b64.json> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    b64_path = sys.argv[2]
    output_path = sys.argv[3]

    for p in [input_path, b64_path]:
        if not os.path.exists(p):
            print(f"ERROR: File not found: {p}")
            sys.exit(1)

    # Load and decode translations
    with open(b64_path, 'r', encoding='utf-8') as f:
        b64_data = json.load(f)

    translations = {}
    for k, v in b64_data.items():
        translations[k] = base64.b64decode(v).decode('utf-8')

    print(f"Loaded {len(translations)} translations")

    # Load document and apply
    doc = Document(input_path)

    applied = 0
    for i, para in enumerate(doc.paragraphs):
        key = str(i)
        if key in translations:
            para.clear()
            para.add_run(translations[key])
            applied += 1

    print(f"Applied {applied} paragraph translations")

    doc.save(output_path)
    print(f"Saved to: {output_path}")


if __name__ == '__main__':
    main()
